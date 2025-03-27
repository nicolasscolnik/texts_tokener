# import os
# import nltk
# nltk.download('punkt_tab')
# import tkinter as tk
# from tkinter import messagebox, ttk
# from collections import Counter
# from nltk.tokenize import word_tokenize, sent_tokenize

# # Intentar importar python-docx y manejar el error sin salir abruptamente
# try:
#     from docx import Document
# except ModuleNotFoundError:
#     Document = None
#     print("Advertencia: Instala 'python-docx' con 'pip install python-docx' para procesar archivos DOCX.")

# # Intentar importar PyPDF2 para PDFs
# try:
#     from PyPDF2 import PdfReader
# except ModuleNotFoundError:
#     PdfReader = None
#     print("Advertencia: Instala 'PyPDF2' con 'pip install PyPDF2' para procesar archivos PDF.")

# # Descargar recursos de NLTK si no están disponibles
# nltk.download('punkt')

# def extract_text_from_docx(docx_path):
#     """Extrae texto de un archivo DOCX."""
#     if Document is None:
#         return "Error: python-docx no está instalado."
#     doc = Document(docx_path)
#     return "\n".join([para.text for para in doc.paragraphs])

# def extract_text_from_pdf(pdf_path):
#     """Extrae texto de un archivo PDF."""
#     if PdfReader is None:
#         return "Error: PyPDF2 no está instalado."
#     with open(pdf_path, "rb") as file:
#         reader = PdfReader(file)
#         text = "\n".join([page.extract_text() for page in reader.pages if page.extract_text()])
#     return text

# def process_text(text):
#     """Procesa el texto y devuelve estadísticas clave."""
#     if not text:
#         return 0, 0, 0, 0, 0, []

#     tokens = word_tokenize(text.lower())
#     tokens = [t for t in tokens if t.isalnum()]

#     num_tokens = len(tokens)
#     unique_tokens = len(set(tokens))
#     num_chars = len(text)
    
#     sentences = sent_tokenize(text)
#     num_sentences = len(sentences)
#     avg_words_per_sentence = num_tokens / num_sentences if num_sentences > 0 else 0

#     freq_dist = Counter(tokens).most_common(10)

#     return num_tokens, unique_tokens, num_chars, num_sentences, avg_words_per_sentence, freq_dist

# def list_files_in_directory(directory):
#     """Lista los archivos disponibles en la carpeta 'textos'."""
#     return [file for file in os.listdir(directory) if file.endswith(('.docx', '.pdf'))]

# def tokenize_selected_file():
#     """Realiza la tokenización del archivo seleccionado."""
#     selected_idx = file_listbox.curselection()
#     if not selected_idx:
#         messagebox.showwarning("Selección requerida", "Por favor, selecciona un archivo.")
#         return
    
#     selected_file = file_listbox.get(selected_idx[0])
#     file_path = os.path.join("textos", selected_file)

#     if selected_file.endswith(".docx"):
#         text = extract_text_from_docx(file_path)
#     elif selected_file.endswith(".pdf"):
#         text = extract_text_from_pdf(file_path)
#     else:
#         messagebox.showerror("Error", "Formato de archivo no soportado.")
#         return

#     if not text:
#         messagebox.showerror("Error", "No se pudo extraer texto del archivo.")
#         return

#     num_tokens, unique_tokens, num_chars, num_sentences, avg_words_per_sentence, freq_dist = process_text(text)

#     # Crear una nueva ventana para los resultados
#     result_window = tk.Toplevel(root)
#     result_window.title("Resultados de Tokenización")
#     result_window.geometry("400x300")

#     # Crear un Frame con ttk para diseño más moderno
#     frame = ttk.Frame(result_window, padding=10)
#     frame.pack(fill="both", expand=True)

#     # Mostrar estadísticas
#     stats_text = (
#         f"Total de palabras (tokens): {num_tokens}\n"
#         f"Palabras únicas (types): {unique_tokens}\n"
#         f"Cantidad de caracteres: {num_chars}\n"
#         f"Cantidad de oraciones: {num_sentences}\n"
#         f"Promedio de palabras por oración: {avg_words_per_sentence:.2f}\n\n"
#         f"Palabras más comunes:\n"
#     )
    
#     stats_label = ttk.Label(frame, text=stats_text, justify="left")
#     stats_label.pack()

#     # Lista de palabras más comunes
#     words_list = ttk.Treeview(frame, columns=("Palabra", "Frecuencia"), show="headings")
#     words_list.heading("Palabra", text="Palabra")
#     words_list.heading("Frecuencia", text="Frecuencia")
#     words_list.column("Palabra", width=150)
#     words_list.column("Frecuencia", width=80)
#     words_list.pack()

#     for word, count in freq_dist:
#         words_list.insert("", tk.END, values=(word, count))

# # Configuración de la ventana principal
# root = tk.Tk()
# root.title("Tokenizador de Documentos")
# root.geometry("500x400")

# # Estilo moderno con ttk
# frame = ttk.Frame(root, padding=10)
# frame.pack(fill="both", expand=True)

# # Etiqueta de instrucción
# ttk.Label(frame, text="Selecciona un archivo para tokenizar:", font=("Arial", 12)).pack(pady=5)

# # Listbox con scrollbar
# listbox_frame = ttk.Frame(frame)
# listbox_frame.pack(fill="both", expand=True)

# file_listbox = tk.Listbox(listbox_frame, width=50, height=10)
# file_listbox.pack(side="left", fill="both", expand=True)

# scrollbar = ttk.Scrollbar(listbox_frame, orient="vertical", command=file_listbox.yview)
# scrollbar.pack(side="right", fill="y")

# file_listbox.config(yscrollcommand=scrollbar.set)

# # Botón para iniciar tokenización
# tokenize_button = ttk.Button(frame, text="Tokenizar", command=tokenize_selected_file)
# tokenize_button.pack(pady=10)

# # Cargar archivos disponibles
# textos_folder = "textos"
# if not os.path.exists(textos_folder):
#     os.makedirs(textos_folder)
# files = list_files_in_directory(textos_folder)
# for file in files:
#     file_listbox.insert(tk.END, file)

# # Iniciar la aplicación
# root.mainloop()

import os
import nltk
import tkinter as tk
from tkinter import messagebox, ttk
from collections import Counter
from nltk.tokenize import word_tokenize, sent_tokenize

# Intentar importar python-docx
try:
    from docx import Document
except ModuleNotFoundError:
    Document = None
    print("Advertencia: Instala 'python-docx' con 'pip install python-docx' para procesar archivos DOCX.")

# Intentar importar PyPDF2
try:
    from PyPDF2 import PdfReader
except ModuleNotFoundError:
    PdfReader = None
    print("Advertencia: Instala 'PyPDF2' con 'pip install PyPDF2' para procesar archivos PDF.")

# Descargar recursos de NLTK si no están disponibles
nltk.download('punkt')

# Obtener ruta absoluta de la carpeta textos
TEXTOS_DIR = os.path.abspath("textos")
if not os.path.exists(TEXTOS_DIR):
    os.makedirs(TEXTOS_DIR)

def list_files_in_directory(directory):
    """Lista los archivos DOCX y PDF dentro del directorio especificado."""
    try:
        files = [
            file for file in os.listdir(directory)
            if file.lower().endswith(('.docx', '.pdf'))
        ]
        print(f"Archivos detectados en {directory}: {files}")  # Depuración
        return files
    except Exception as e:
        print(f"Error al listar archivos: {e}")
        return []

def extract_text_from_docx(docx_path):
    """Extrae texto de un archivo DOCX."""
    if Document is None:
        return "Error: python-docx no está instalado."
    try:
        doc = Document(docx_path)
        return "\n".join([para.text for para in doc.paragraphs])
    except Exception as e:
        return f"Error al leer DOCX: {e}"

def extract_text_from_pdf(pdf_path):
    """Extrae texto de un archivo PDF."""
    if PdfReader is None:
        return "Error: PyPDF2 no está instalado."
    try:
        with open(pdf_path, "rb") as file:
            reader = PdfReader(file)
            text = "\n".join([page.extract_text() for page in reader.pages if page.extract_text()])
        return text
    except Exception as e:
        return f"Error al leer PDF: {e}"

def process_text(text):
    """Procesa el texto y devuelve estadísticas clave."""
    if not text:
        return 0, 0, 0, 0, 0, []

    tokens = word_tokenize(text.lower())
    tokens = [t for t in tokens if t.isalnum()]

    num_tokens = len(tokens)
    unique_tokens = len(set(tokens))
    num_chars = len(text)
    
    sentences = sent_tokenize(text)
    num_sentences = len(sentences)
    avg_words_per_sentence = num_tokens / num_sentences if num_sentences > 0 else 0

    freq_dist = Counter(tokens).most_common(10)

    return num_tokens, unique_tokens, num_chars, num_sentences, avg_words_per_sentence, freq_dist

def tokenize_selected_file():
    """Realiza la tokenización del archivo seleccionado."""
    selected_idx = file_listbox.curselection()
    if not selected_idx:
        messagebox.showwarning("Selección requerida", "Por favor, selecciona un archivo.")
        return
    
    selected_file = file_listbox.get(selected_idx[0])
    file_path = os.path.join(TEXTOS_DIR, selected_file)

    if selected_file.lower().endswith(".docx"):
        text = extract_text_from_docx(file_path)
    elif selected_file.lower().endswith(".pdf"):
        text = extract_text_from_pdf(file_path)
    else:
        messagebox.showerror("Error", "Formato de archivo no soportado.")
        return

    if not text or text.startswith("Error"):
        messagebox.showerror("Error", f"No se pudo extraer texto del archivo.\n{text}")
        return

    num_tokens, unique_tokens, num_chars, num_sentences, avg_words_per_sentence, freq_dist = process_text(text)

    # Crear ventana de resultados
    result_window = tk.Toplevel(root)
    result_window.title("Resultados de Tokenización")
    result_window.geometry("400x300")

    frame = ttk.Frame(result_window, padding=10)
    frame.pack(fill="both", expand=True)

    stats_text = (
        f"Total de palabras (tokens): {num_tokens}\n"
        f"Palabras únicas (types): {unique_tokens}\n"
        f"Cantidad de caracteres: {num_chars}\n"
        f"Cantidad de oraciones: {num_sentences}\n"
        f"Promedio de palabras por oración: {avg_words_per_sentence:.2f}\n\n"
        f"Palabras más comunes:\n"
    )
    
    stats_label = ttk.Label(frame, text=stats_text, justify="left")
    stats_label.pack()

    words_list = ttk.Treeview(frame, columns=("Palabra", "Frecuencia"), show="headings")
    words_list.heading("Palabra", text="Palabra")
    words_list.heading("Frecuencia", text="Frecuencia")
    words_list.column("Palabra", width=150)
    words_list.column("Frecuencia", width=80)
    words_list.pack()

    for word, count in freq_dist:
        words_list.insert("", tk.END, values=(word, count))

# Configuración de la ventana principal
root = tk.Tk()
root.title("Tokenizador de Documentos")
root.geometry("500x400")

frame = ttk.Frame(root, padding=10)
frame.pack(fill="both", expand=True)

ttk.Label(frame, text="Selecciona un archivo para tokenizar:", font=("Arial", 12)).pack(pady=5)

listbox_frame = ttk.Frame(frame)
listbox_frame.pack(fill="both", expand=True)

file_listbox = tk.Listbox(listbox_frame, width=50, height=10)
file_listbox.pack(side="left", fill="both", expand=True)

scrollbar = ttk.Scrollbar(listbox_frame, orient="vertical", command=file_listbox.yview)
scrollbar.pack(side="right", fill="y")

file_listbox.config(yscrollcommand=scrollbar.set)

tokenize_button = ttk.Button(frame, text="Tokenizar", command=tokenize_selected_file)
tokenize_button.pack(pady=10)

# Cargar archivos disponibles
files = list_files_in_directory(TEXTOS_DIR)
for file in files:
    file_listbox.insert(tk.END, file)

root.mainloop()

